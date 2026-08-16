import os
from docx import Document

# Dictionary of replacements (Old Text : New Text)
REPLACEMENTS = {
    "[Project Title]": "FarmXP",
    "[SIH25XXX]": "FSJ28-INTERN-046",
    "[Gamified Platform to Promote Sustainable Farming Practices]": "FarmXP - Gamified Platform to Promote Sustainable Farming Practices",
    "[FSJ28-INTERN-046]": "FSJ28-INTERN-046",
    "[ ] Software only": "[x] Software only",
    "[✅] Software only": "[x] Software only",
    "[ Add project-specific terms here ]": "XP: Experience points. Leaderboard: Rank of farmers.",
    "[ Add: any dataset sources, API references, standards specific to your domain ]": "Google Gemini API Documentation",
    "[ Role 4 — e.g., Reviewer / Approver ]": "Not Applicable",
    "[ Role 5 — project-specific ]": "Not Applicable",
    "[ local machine / AWS EC2 / Kubernetes (minikube) / Railway / Render ]": "local machine / Docker",
    "MySQL (relational) · MongoDB 6.0+ (document/time-series)": "Oracle Database (Relational)",
    "MySQL 8.0": "Oracle Database",
    "MongoDB 6.0": "Not Applicable",
    "MySQL/MongoDB": "Oracle Database",
    "[ Arduino Uno / Raspberry Pi 4 / ESP32 ] + [ Sensor types and models ]": "Not Applicable (Software Only)",
    "[ WiFi / 4G / LoRa ] for IoT data (hybrid)": "4G/WiFi",
    "(total hardware budget < ₹2,000) ] (Hybrid only)": "(total hardware budget < Rs 2,000) ] (Not Applicable)",
    "[ Add project-specific constraints: API rate limits, data privacy regulations, language requirements ]": "Gemini API rate limits (15 RPM for free tier).",
    "[ Add project-specific assumption ]": "Farmers will log accurate practices for verification.",
    "[ Specific government data API — e.g., Agmarknet, CPCB sensor feed, CGWB DWLR data ]": "None currently integrated.",
    "[ Satellite/geospatial data source — e.g., Sentinel Hub, NASA POWER, ISRO Bhuvan ]": "None currently integrated.",
    "[ Add any hardware component dependency ]": "Not Applicable",
    "[ Service 1 — e.g., UserService ]": "auth-service",
    "[ Service 2 — e.g., CoreDomainService ]": "farmer-service",
    "[ Primary domain logic — crops, patients, assets, etc. ]": "Farmer profile management, XP calculation",
    "[ Service 3 — e.g., AnalyticsService ]": "learning-service",
    "[ Analytics, reporting, aggregated data views ]": "Learning modules, quizzes, progress tracking",
    "[ Service 4 — e.g., NotificationService ]": "notification-service",
    "[ IoT Ingestion Service (Hybrid) ]": "sustainability-service",
    "Receives and stores IoT sensor telemetry via REST/MQTT": "Calculates sustainability score and tracks certified practices",
    "[ Microcontroller ]": "Not Applicable",
    "[ Sensor 1 ]": "Not Applicable",
    "[ e.g., DHT22 Temperature/Humidity ]": "Not Applicable",
    "[ Measure ambient conditions ]": "Not Applicable",
    "[ Sensor 2 ]": "Not Applicable",
    "[ e.g., MQ-135 Gas Sensor ]": "Not Applicable",
    "[ Detect specific gas concentrations ]": "Not Applicable",
    "[ Sensor 3 ]": "Not Applicable",
    "[ e.g., HC-SR04 Ultrasonic ]": "Not Applicable",
    "[ Distance / level measurement ]": "Not Applicable",
    "[ Communication Module ]": "Not Applicable",
    "[ GSM SIM800 / LoRa SX1278 / WiFi (ESP8266) ]": "Not Applicable",
    "[ Power Supply ]": "Not Applicable",
    "[ Screen 3 — core user action ]": "Learning Portal",
    "[ Primary user ]": "Farmer",
    "[ Describe form, map, chart, or data entry screen ]": "View available modules, complete quizzes.",
    "[ Screen 4 — AI feature screen ]": "AI Assistant",
    "[ User triggering AI ]": "Farmer",
    "[ Upload photo / enter query / view AI result — describe input and output display ]": "Chat interface to ask for recommendations.",
    "[ Screen 5 — admin / officer view ]": "Admin Dashboard",
    "[ Screen 6 — reports ]": "Scheme Management",
    "[ e.g., Soil Moisture Sensor ]": "Not Applicable",
    "[ e.g., Gas Sensor ]": "Not Applicable",
    "OpenAI API (GPT-4o)": "Google Gemini API",
    "GenAI features — natural language responses, report generation, chatbot": "Chatbot recommendations based on farmer profile.",
    "HuggingFace Inference API": "Not Applicable",
    "NLP model inference (BERT, sentence-transformers) without local GPU": "Not Applicable",
    "[ Domain-specific API ]": "Not Applicable",
    "[ e.g., Agmarknet / CGWB / CPCB ]": "Not Applicable",
    "[ API Key / OAuth ]": "Not Applicable",
    "[ Describe specific use ]": "Not Applicable",
    "[ Describe in 1-2 sentences what this AI module does and what human effort it replaces or augments ]": "Replaces manual browsing by automatically analyzing farmer XP, score, and modules to recommend the next best module.",
    "JPG image of crop leaf, 224x224 pixels": "Structured JSON with Farmer Context (XP, Sustainability Score, Completed Modules) and available Module titles/categories.",
    "Disease class (one of 10) with confidence score 0-1": "Natural language recommendation and advice.",
    "MobileNetV2 CNN / XGBoost / BERT fine-tuned / OpenAI GPT-4o API / Isolation Forest / LSTM": "Google Gemini LLM via Spring AI",
    "Dataset name + source URL": "Not Applicable (Pre-trained LLM)",
    "[ HTTP method ] /[ endpoint path ]": "POST /api/ai/chat",
    "[SpringBoot ServiceName]": "ai-service",
    "Classification accuracy > 80% / RMSE < 15% / Precision > 75% on test set": "Response latency < 10s, highly relevant recommendations.",
    "Display manual review flag to field officer": "Returns a fallback message stating the AI service is busy (rate limit handling).",
    "AI output is advisory only; human officer makes final decision": "AI output is purely advisory for learning purposes.",
    "[ Service 2 Name ]": "learning-service",
    "GET/[ resource ]/{id}": "GET /api/farmers/profile",
    "POST/[ resource ]": "POST /api/sustainability/practices",
    "PUT/[ resource ]/{id}": "PUT /api/admin/practices/{id}/approve",
    "DELETE/[ resource ]/{id}": "DELETE /api/notifications/{id}",
    "GET/[ resource ]/[ filter ]": "GET /api/market/matches",
    "GET/[ resource ]": "GET /api/learning/modules",
    "SUSTAINABILITY_MATRIC": "SUSTAINABILITY_METRIC",
    "VARCAHR": "VARCHAR",
    "Project Title": "FarmXP",
    "Gamified Platform to Promote Sustainable Farming Practices": "FarmXP - Gamified Platform to Promote Sustainable Farming Practices"
}

def replace_text_in_paragraph(paragraph, replacements):
    """
    Replaces text in a paragraph while attempting to preserve formatting runs.
    If a placeholder is split across multiple runs, it merges the text into the first run
    and clears the subsequent runs to maintain document structure.
    """
    # If no replacement text is in this paragraph, skip to save time
    if not any(key in paragraph.text for key in replacements):
        return

    for key, val in replacements.items():
        if key in paragraph.text:
            # Simple case: the key exists entirely within a single run
            for run in paragraph.runs:
                if key in run.text:
                    run.text = run.text.replace(key, val)
            
            # Complex case: the key is split across multiple runs
            if key in paragraph.text:
                # To perfectly preserve formatting of the start of the string,
                # we combine all runs, do the replacement, and put the text back 
                # into the first run, clearing the others.
                # This keeps paragraph-level styles completely intact.
                full_text = "".join(run.text for run in paragraph.runs)
                full_text = full_text.replace(key, val)
                
                if len(paragraph.runs) > 0:
                    paragraph.runs[0].text = full_text
                    for i in range(1, len(paragraph.runs)):
                        paragraph.runs[i].text = ""

def process_document(input_path, output_path):
    print(f"Opening template: {input_path}")
    doc = Document(input_path)

    # 1. Replace in standard paragraphs
    for para in doc.paragraphs:
        replace_text_in_paragraph(para, REPLACEMENTS)

    # 2. Replace in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_text_in_paragraph(para, REPLACEMENTS)
                    
    # 3. Replace in headers and footers
    for section in doc.sections:
        for header in [section.header, section.first_page_header, section.even_page_header]:
            if header:
                for para in header.paragraphs:
                    replace_text_in_paragraph(para, REPLACEMENTS)
                for table in header.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                replace_text_in_paragraph(para, REPLACEMENTS)
                                
        for footer in [section.footer, section.first_page_footer, section.even_page_footer]:
            if footer:
                for para in footer.paragraphs:
                    replace_text_in_paragraph(para, REPLACEMENTS)
                for table in footer.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                replace_text_in_paragraph(para, REPLACEMENTS)

    print(f"Saving completed document to: {output_path}")
    doc.save(output_path)
    print("Done! Open the output file to verify.")

if __name__ == "__main__":
    TEMPLATE = r"D:\FULL STACK JAVA\FarmXp - Integration\SRS_FSJ28-INTERN-046_Tech-Titans_v1.0.docx"
    OUTPUT = r"D:\FULL STACK JAVA\FarmXp - Integration\FarmXP_Completed_Form.docx"
    
    if not os.path.exists(TEMPLATE):
        print(f"Error: Could not find template file at {TEMPLATE}")
    else:
        process_document(TEMPLATE, OUTPUT)
